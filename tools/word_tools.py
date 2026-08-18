"""
Word 文档读写工具
支持 .docx 格式的读取和生成
"""

import io
import json
from docx import Document
from .json_repair import safe_parse_json
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


def _extract_paragraph_text(doc: Document) -> list[dict]:
    """提取文档中所有段落的文本和样式"""
    result = []
    for i, p in enumerate(doc.paragraphs):
        if not p.text.strip():
            continue
        runs_info = []
        for r in p.runs:
            runs_info.append({
                "text": r.text,
                "bold": r.bold,
                "font_name": r.font.name,
                "font_size": str(r.font.size) if r.font.size else None,
            })
        result.append({
            "index": i,
            "text": p.text,
            "alignment": str(p.alignment),
            "style": p.style.name if p.style else None,
            "runs": runs_info,
        })
    return result


def _extract_tables(doc: Document) -> list[dict]:
    """提取文档中所有表格"""
    result = []
    for ti, t in enumerate(doc.tables):
        rows_data = []
        for ri, row in enumerate(t.rows):
            cells = [cell.text for cell in row.cells]
            rows_data.append(cells)
        result.append({
            "table_index": ti,
            "rows": len(t.rows),
            "cols": len(t.columns),
            "data": rows_data,
        })
    return result


def read_docx(file_path: str, mode: str = "full") -> str:
    """
    读取 Word 文档内容

    Args:
        file_path: .docx 文件路径
        mode: 读取模式
              - "full": 返回段落文字和表格（默认）
              - "paragraphs": 仅段落文字
              - "tables": 仅表格内容
              - "structure": 返回段落的样式和格式信息

    Returns:
        JSON 格式的文档内容字符串
    """
    try:
        doc = Document(file_path)
    except Exception as e:
        return json.dumps({"error": f"无法打开文件: {str(e)}"}, ensure_ascii=False)

    output = {"file": file_path, "mode": mode}

    if mode in ("full", "paragraphs", "structure"):
        paragraphs = _extract_paragraph_text(doc)
        if mode == "structure":
            output["paragraphs"] = paragraphs
        else:
            output["paragraphs"] = [p["text"] for p in paragraphs]

    if mode in ("full", "tables"):
        output["tables"] = _extract_tables(doc)

    return json.dumps(output, ensure_ascii=False, indent=2)


# ═══════════════════════════════════════════════════
# 写入功能
# ═══════════════════════════════════════════════════

def _set_cn_font(run, font_name="宋体", size=None, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold


def _add_paragraph(doc, text, font_name="宋体", size=12, bold=False,
                   alignment=None, space_after=0, first_line_indent=None):
    """添加段落并设置格式"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    _set_cn_font(run, font_name, size, bold)
    return p


def _shade_cell(cell, color):
    """设置单元格背景"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _vertical_center(cell):
    """单元格垂直居中"""
    tcPr = cell._tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="center"/>')
    tcPr.append(vAlign)


def _make_table(doc, headers, rows, col_widths=None,
                header_bg="333333", header_fg="FFFFFF",
                summary_row_bg="F2F2F2", font_size=10.5):
    """
    创建格式化表格

    Args:
        doc: Document 对象
        headers: 表头列表
        rows: 数据行列表（二维列表）
        col_widths: 列宽列表（Cm），可选
        header_bg: 表头背景色
        header_fg: 表头字体色
        summary_row_bg: 汇总行背景色
        font_size: 字体大小
    """
    n_rows = len(rows) + 1
    n_cols = len(headers)
    table = doc.add_table(rows=n_rows, cols=n_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    # 列宽
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = w

    # 表头
    for i, h in enumerate(headers):
        cell = table.cell(0, i)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        _set_cn_font(run, "宋体", font_size, bold=True)
        run.font.color.rgb = RGBColor(
            int(header_fg[0:2], 16),
            int(header_fg[2:4], 16),
            int(header_fg[4:6], 16),
        )
        _shade_cell(cell, header_bg)
        _vertical_center(cell)

    # 数据行
    for ri, row_data in enumerate(rows):
        is_last = (ri == len(rows) - 1)
        for ci, val in enumerate(row_data):
            cell = table.cell(ri + 1, ci)
            p = cell.paragraphs[0]
            if ci >= len(headers) - 2:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            elif ci == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(str(val))
            _set_cn_font(run, "宋体", font_size, bold=is_last)
            _vertical_center(cell)
            if is_last:
                _shade_cell(cell, summary_row_bg)

    return table


def write_docx(spec_json: str) -> str:
    """
    根据 JSON 规格生成 .docx 文件

    spec_json 格式:
    {
        "output": "/path/to/output.docx",       // 必填，输出路径
        "title": "合同标题",                     // 可选，文档标题
        "title_font": "黑体",                    // 标题字体，默认黑体
        "title_size": 22,                        // 标题字号
        "body_font": "宋体",                     // 正文字体
        "body_size": 12,                         // 正文字号
        "heading_font": "黑体",                  // 条款标题字体
        "heading_size": 14,                      // 条款标题字号
        "content": [
            {"type": "heading", "text": "第一条  项目概况", "level": 1},
            {"type": "paragraph", "text": "正文内容...", "indent": true},
            {"type": "empty"},
            {"type": "table", "headers": ["序号","名称","数量"], "rows": [["1","交换机","3"]], "widths": [1.5, 5.5, 2.0]},
            {"type": "page_break"},
            {"type": "signature", "left": {"party":"甲方","rep":"______"}, "right": {"party":"乙方","rep":"张三"}}
        ]
    }
    """
    spec, err = safe_parse_json(spec_json)
    if err:
        return json.dumps({"error": f"JSON 解析失败: {err}"}, ensure_ascii=False)

    output_path = spec.get("output")
    if not output_path:
        return json.dumps({"error": "必须指定 output 路径"}, ensure_ascii=False)

    doc = Document()

    # 页面设置
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # 全局样式
    style = doc.styles['Normal']
    style.font.name = spec.get("body_font", "宋体")
    style.font.size = Pt(spec.get("body_size", 12))
    style.element.rPr.rFonts.set(qn('w:eastAsia'), spec.get("body_font", "宋体"))
    style.paragraph_format.line_spacing = 1.5

    title_font = spec.get("title_font", "黑体")
    title_size = spec.get("title_size", 22)
    body_font = spec.get("body_font", "宋体")
    body_size = spec.get("body_size", 12)
    heading_font = spec.get("heading_font", "黑体")
    heading_size = spec.get("heading_size", 14)

    # 渲染内容
    for item in spec.get("content", []):
        t = item.get("type", "paragraph")

        if t == "title":
            _add_paragraph(doc, "", size=body_size)
            for line in item.get("text", "").split("\n"):
                _add_paragraph(doc, line, font_name=title_font, size=title_size,
                              bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(doc, "", size=body_size)

        elif t == "heading":
            _add_paragraph(doc, "", size=6)
            _add_paragraph(doc, item.get("text", ""), font_name=heading_font,
                          size=heading_size, bold=True)

        elif t == "paragraph":
            indent = 0.74 if item.get("indent") else None
            _add_paragraph(doc, item.get("text", ""), font_name=body_font,
                          size=body_size, first_line_indent=indent)

        elif t == "empty":
            _add_paragraph(doc, "", size=body_size)

        elif t == "table":
            _make_table(
                doc,
                headers=item.get("headers", []),
                rows=item.get("rows", []),
                col_widths=[Cm(w) for w in item.get("widths", [])] if item.get("widths") else None,
                font_size=item.get("font_size", 10.5),
            )

        elif t == "page_break":
            doc.add_page_break()

        elif t == "signature":
            _add_paragraph(doc, "", size=12)
            _add_paragraph(doc, "（本页为签署页，无正文）", font_name=body_font,
                          size=10.5, alignment=WD_ALIGN_PARAGRAPH.CENTER)
            _add_paragraph(doc, "", size=24)

            left = item.get("left", {})
            right = item.get("right", {})

            sign_table = doc.add_table(rows=6, cols=2)
            sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            sign_data = [
                (f'{left.get("party", "甲方")}（盖章）', f'{right.get("party", "乙方")}（盖章）'),
                ("", right.get("company", "")),
                ("", ""),
                ("授权代表签字", "授权代表签字"),
                (left.get("rep", "__________________"), right.get("rep", "__________________")),
                (f'日期：{left.get("date", "______年____月____日")}',
                 f'日期：{right.get("date", "______年____月____日")}'),
            ]
            for ri, (l, r) in enumerate(sign_data):
                for ci, text in enumerate([l, r]):
                    cell = sign_table.cell(ri, ci)
                    p = cell.paragraphs[0]
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.paragraph_format.line_spacing = 1.8
                    run = p.add_run(text)
                    _set_cn_font(run, body_font, 12, bold=(ri in [0, 3]))
                    _vertical_center(cell)
            sign_table.cell(0, 0).width = Cm(7)
            sign_table.cell(0, 1).width = Cm(7)

    # 页码
    for sec in doc.sections:
        footer = sec.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
        run2 = p.add_run()
        run2._r.append(parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>'))
        run3 = p.add_run()
        run3._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))

    doc.save(output_path)
    return json.dumps({"success": True, "output": output_path}, ensure_ascii=False)


# ═══════════════════════════════════════════════════
# 编辑功能（修改已有 .docx 文件）
# ═══════════════════════════════════════════════════

def _replace_in_paragraph(p, find, replace):
    """在段落内查找替换文本，保留首个 run 的格式"""
    if find not in p.text:
        return False
    new_text = p.text.replace(find, replace)
    if p.runs:
        p.runs[0].text = new_text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(new_text)
    return True


def _set_paragraph_text(p, text):
    """替换整个段落文本（保留首个 run 格式）"""
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def _delete_paragraph(p):
    """删除段落"""
    p._element.getparent().remove(p._element)


def _check_para_index(op: dict, n_para: int, allow_append: bool = False):
    """校验操作中的段落 index（相对原始文档）；返回错误 JSON 字符串或 None"""
    idx = op.get("index", 0)
    if isinstance(idx, bool) or not isinstance(idx, int):
        return json.dumps({"error": f"index 必须是整数: {idx!r}"}, ensure_ascii=False)
    upper = n_para + (1 if allow_append else 0)
    if not (0 <= idx < upper):
        return json.dumps({
            "error": f"index {idx} 超出范围（原始文档共 {n_para} 段，有效值 0-{upper - 1}）",
        }, ensure_ascii=False)
    return None


def _para_alive(p) -> bool:
    """段落是否仍在文档树中（未被之前的操作删除）"""
    return p._element.getparent() is not None


def edit_docx(file_path: str, spec_json: str) -> str:
    """
    编辑已有 Word 文档（.docx）

    spec_json 格式:
    {
        "output": "/path/to/output.docx",     // 可选，不指定则覆盖原文件
        "operations": [
            {"op": "replace_all", "find": "旧文字", "replace": "新文字"},             // 全局查找替换（段落+表格）
            {"op": "insert_paragraph", "index": 3, "text": "新段落", "bold": false},  // 在第 index 段之前插入
            {"op": "delete_paragraph", "index": 5},                                   // 删除指定段落
            {"op": "set_paragraph_text", "index": 2, "text": "新文本"},                // 替换指定段落文本
            {"op": "add_paragraph", "text": "末尾追加", "bold": true},                 // 文档末尾追加段落
            {"op": "set_table_cell", "table_index": 0, "row": 1, "col": 0, "text": "新值"}  // 修改表格单元格
        ]
    }

    注意：所有段落 index 均相对于文档初始状态（与 read_docx 返回的编号一致），
    不受本批次中前面的插入/删除操作影响。
    """
    spec, err = safe_parse_json(spec_json)
    if err:
        return json.dumps({"error": f"JSON 解析失败: {err}"}, ensure_ascii=False)

    try:
        doc = Document(file_path)
    except Exception as e:
        return json.dumps({"error": f"无法打开文件: {str(e)}"}, ensure_ascii=False)

    ops = spec.get("operations", [])
    if not ops:
        return json.dumps({"error": "operations 不能为空"}, ensure_ascii=False)

    # 原始段落快照：所有 index 均相对于文档初始状态（与 read_docx 编号一致），
    # 不受本批次中前面的插入/删除操作影响
    original_paras = list(doc.paragraphs)
    n_para = len(original_paras)

    for op in ops:
        kind = op.get("op")
        try:
            if kind == "replace_all":
                find, replace = op.get("find", ""), op.get("replace", "")
                for p in doc.paragraphs:
                    _replace_in_paragraph(p, find, replace)
                for t in doc.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            for p in cell.paragraphs:
                                _replace_in_paragraph(p, find, replace)

            elif kind == "insert_paragraph":
                err = _check_para_index(op, n_para, allow_append=True)
                if err:
                    return err
                idx = op.get("index", 0)
                if idx < n_para:
                    target = original_paras[idx]
                    if not _para_alive(target):
                        return json.dumps({
                            "error": f"段落 index={idx} 已被之前的操作删除，无法在其前插入",
                        }, ensure_ascii=False)
                    p = target.insert_paragraph_before()
                else:
                    p = doc.add_paragraph()
                run = p.add_run(op.get("text", ""))
                _set_cn_font(run, "宋体", 12, bold=op.get("bold", False))

            elif kind == "delete_paragraph":
                err = _check_para_index(op, n_para)
                if err:
                    return err
                target = original_paras[op.get("index", 0)]
                if not _para_alive(target):
                    return json.dumps({
                        "error": f"段落 index={op.get('index', 0)} 已被之前的操作删除",
                    }, ensure_ascii=False)
                _delete_paragraph(target)

            elif kind == "set_paragraph_text":
                err = _check_para_index(op, n_para)
                if err:
                    return err
                target = original_paras[op.get("index", 0)]
                if not _para_alive(target):
                    return json.dumps({
                        "error": f"段落 index={op.get('index', 0)} 已被之前的操作删除",
                    }, ensure_ascii=False)
                _set_paragraph_text(target, op.get("text", ""))

            elif kind == "add_paragraph":
                _add_paragraph(doc, op.get("text", ""), bold=op.get("bold", False))

            elif kind == "set_table_cell":
                ti = op.get("table_index", 0)
                if ti < len(doc.tables):
                    table = doc.tables[ti]
                    r, c = op.get("row", 0), op.get("col", 0)
                    if r < len(table.rows) and c < len(table.columns):
                        table.rows[r].cells[c].text = op.get("text", "")

            else:
                return json.dumps({"error": f"未知操作: {kind}"}, ensure_ascii=False)

        except Exception as e:
            return json.dumps({"error": f"操作 '{kind}' 执行失败: {str(e)}"}, ensure_ascii=False)

    output_path = spec.get("output", file_path)
    doc.save(output_path)
    return json.dumps({"success": True, "output": output_path}, ensure_ascii=False)
