#!/usr/bin/env python3
"""生成规范格式的芜湖古城物业施工合同"""

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import copy

doc = Document()

# ── 页面设置 ──
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.top_margin = Cm(2.54)
section.bottom_margin = Cm(2.54)
section.left_margin = Cm(3.18)
section.right_margin = Cm(3.18)

# ── 样式定义 ──
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(0)

def set_cn_font(run, font_name='宋体', size=None, bold=False):
    """设置中文字体"""
    run.font.name = font_name
    run.element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    if size:
        run.font.size = Pt(size)
    run.bold = bold

def add_paragraph_cn(doc, text, font_name='宋体', size=12, bold=False, alignment=None, space_after=0, first_line_indent=None):
    """添加中文段落"""
    p = doc.add_paragraph()
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.5
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    run = p.add_run(text)
    set_cn_font(run, font_name, size, bold)
    return p

def set_cell_border(cell, **kwargs):
    """设置单元格边框"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
    for edge, val in kwargs.items():
        element = parse_xml(
            f'<w:{edge} {nsdecls("w")} w:val="{val.get("val", "single")}" '
            f'w:sz="{val.get("sz", 4)}" w:space="0" w:color="{val.get("color", "000000")}"/>'
        )
        tcBorders.append(element)
    tcPr.append(tcBorders)

def shade_cell(cell, color):
    """设置单元格背景色"""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)

def set_cell_vertical_alignment(cell, align="center"):
    """设置单元格垂直对齐"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    vAlign = parse_xml(f'<w:vAlign {nsdecls("w")} w:val="{align}"/>')
    tcPr.append(vAlign)

# ═══════════════════════════════════════════════
# 标题
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=12)  # 空行
add_paragraph_cn(doc, '芜湖古城物业停车场系统和监控系统', font_name='黑体', size=22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph_cn(doc, '网络分离项目施工合同', font_name='黑体', size=22, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph_cn(doc, '', size=12)

# ═══════════════════════════════════════════════
# 合同编号和签订信息
# ═══════════════════════════════════════════════
info_table = doc.add_table(rows=6, cols=4)
info_table.alignment = WD_TABLE_ALIGNMENT.CENTER

# 隐藏表格边框（无边框样式）
for row in info_table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            paragraph.paragraph_format.space_after = Pt(2)
            paragraph.paragraph_format.space_before = Pt(2)
            paragraph.paragraph_format.line_spacing = 1.5

# 甲方信息
# 填充甲乙方信息表格
party_labels = [
    ('甲方（委托方）', '', '乙方（承建方）', '________________________'),
    ('统一社会信用代码', '________________________', '统一社会信用代码', '________________________'),
    ('联系地址', '________________________', '联系地址', '________________________'),
    ('法定代表人', '________________________', '法定代表人', '________________________'),
    ('联系人', '________________________', '联系人', '________________________'),
    ('联系电话', '________________________', '联系电话', '________________________'),
]

for row_idx, (a_label, a_val, b_label, b_val) in enumerate(party_labels):
    is_first = (row_idx == 0)
    # 甲方 - 标签列
    cell = info_table.cell(row_idx, 0)
    p = cell.paragraphs[0]
    run = p.add_run(a_label)
    set_cn_font(run, '宋体', 12, bold=is_first)
    p.paragraph_format.line_spacing = 1.5
    if is_first:
        p.paragraph_format.space_before = Pt(6)
    # 甲方 - 值列
    cell = info_table.cell(row_idx, 1)
    p = cell.paragraphs[0]
    run = p.add_run(f'：{a_val}' if a_val else '：________________________')
    set_cn_font(run, '宋体', 12)
    p.paragraph_format.line_spacing = 1.5
    if is_first:
        p.paragraph_format.space_before = Pt(6)
    # 乙方 - 标签列
    cell = info_table.cell(row_idx, 2)
    p = cell.paragraphs[0]
    run = p.add_run(b_label)
    set_cn_font(run, '宋体', 12, bold=is_first)
    p.paragraph_format.line_spacing = 1.5
    if is_first:
        p.paragraph_format.space_before = Pt(6)
    # 乙方 - 值列
    cell = info_table.cell(row_idx, 3)
    p = cell.paragraphs[0]
    run = p.add_run(f'：{b_val}' if b_val else '：________________________')
    set_cn_font(run, '宋体', 12)
    p.paragraph_format.line_spacing = 1.5
    if is_first:
        p.paragraph_format.space_before = Pt(6)

# 调整列宽
for row in info_table.rows:
    row.cells[0].width = Cm(3.2)
    row.cells[1].width = Cm(3.8)
    row.cells[2].width = Cm(3.2)
    row.cells[3].width = Cm(3.8)

# 移除表格默认边框
for row in info_table.rows:
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            f'<w:tcBorders {nsdecls("w")}>'
            f'<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
            f'</w:tcBorders>'
        )
        tcPr.append(tcBorders)

add_paragraph_cn(doc, '', size=12)

# ═══════════════════════════════════════════════
# 鉴于条款
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '鉴于：', font_name='黑体', size=12, bold=True, first_line_indent=0.74)
add_paragraph_cn(doc,
    '根据《中华人民共和国民法典》及相关法律法规，甲乙双方本着平等、自愿、公平、诚实信用的原则，'
    '就本项目施工事宜协商一致，订立本合同。',
    first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第一条 项目概况与施工方案
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第一条  项目概况与施工方案', font_name='黑体', size=14, bold=True)

add_paragraph_cn(doc, '1.1  项目名称：芜湖古城物业停车场系统和监控系统网络分离项目。', first_line_indent=0.74)

add_paragraph_cn(doc, '1.2  实施范围：对现有物业停车场系统与监控系统的网络设备进行物理/逻辑隔离改造，'
    '部署专用交换设备，完成点位布线、设备安装、系统配置、联调测试及验收前现场清理。', first_line_indent=0.74)

add_paragraph_cn(doc, '1.3  施工方案：', first_line_indent=0.74)
add_paragraph_cn(doc, '（1）乙方负责采购本合同附件一清单所列全部设备及辅材；', first_line_indent=0.74)
add_paragraph_cn(doc, '（2）施工采用"先备后切"方案，保障物业业务连续性；新旧网络割接期间，双方共同制定应急预案并提前报备；', first_line_indent=0.74)
add_paragraph_cn(doc, '（3）设备安装位置由甲方指定机房/弱电井确定，乙方负责上架、标签标识、绝缘测试及系统策略配置；', first_line_indent=0.74)
add_paragraph_cn(doc, '（4）调试完成后，双方按本合同第四条进行整体验收。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第二条 合同工期
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第二条  合同工期', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '2.1  施工周期：7个工作日（自合同签订且首笔款项到达乙方指定账户之日起算）。', first_line_indent=0.74)
add_paragraph_cn(doc, '2.2  如遇不可抗力或甲方原因导致无法施工，工期相应顺延；因乙方原因延误的，按第七条承担违约责任。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第三条 合同价款与发票
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第三条  合同价款与发票', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '3.1  本合同含税总价为：人民币（大写）叁仟玖佰玖拾伍元整（¥3,995.00）。', first_line_indent=0.74)
add_paragraph_cn(doc, '3.2  该费用包含设备费、运输费、安装人工费、调试费、税费及质保期内常规服务费等全部履约成本，'
    '除本合同明确约定外不再计取其他任何费用。', first_line_indent=0.74)
add_paragraph_cn(doc, '3.3  乙方应于项目验收合格后【7】个工作日内向甲方开具 1% 增值税专用发票，发票信息以甲方提供为准。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第四条 付款方式
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第四条  付款方式', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '4.1  合同签订后【3】个工作日内，甲方向乙方支付合同总价的 100% 作为工程预付款/进度款；'
    '项目竣工验收合格且收到合规发票后【7】个工作日内，双方完成财务对账（如已全额预付则本条不适用）。', first_line_indent=0.74)
add_paragraph_cn(doc, '4.2  乙方收款账户信息：', first_line_indent=0.74)

# 银行账户信息
bank_table = doc.add_table(rows=3, cols=2)
bank_table.alignment = WD_TABLE_ALIGNMENT.LEFT
bank_data = [
    ('户名', '________________________'),
    ('开户行', '________________________'),
    ('账号', '________________________'),
]
for i, (label, value) in enumerate(bank_data):
    c0 = bank_table.cell(i, 0)
    c1 = bank_table.cell(i, 1)
    p0 = c0.paragraphs[0]
    r0 = p0.add_run(f'    {label}：')
    set_cn_font(r0, '宋体', 12, bold=True)
    p0.paragraph_format.line_spacing = 1.5
    p1 = c1.paragraphs[0]
    r1 = p1.add_run(value)
    set_cn_font(r1, '宋体', 12)
    p1.paragraph_format.line_spacing = 1.5
    # Remove borders
    for c in [c0, c1]:
        tcPr = c._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)

bank_table.cell(0, 0).width = Cm(2.5)
bank_table.cell(0, 1).width = Cm(9)

# ═══════════════════════════════════════════════
# 第五条 质保期限与服务承诺
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第五条  质保期限与服务承诺', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '5.1  质保期：自项目整体验收合格之日起 壹年（12个月）。', first_line_indent=0.74)
add_paragraph_cn(doc, '5.2  质保范围：清单内硬件设备因非人为损坏、非不可抗力导致的故障，乙方提供免费维修或更换；'
    '提供系统配置支持及网络策略优化服务。', first_line_indent=0.74)
add_paragraph_cn(doc, '5.3  响应时效：乙方提供 7×24 小时技术支持，接到报修后 2 小时内响应，'
    '如需现场处理，24 小时内到达指定地点。', first_line_indent=0.74)
add_paragraph_cn(doc, '5.4  质保期满后，双方可另行协商维保协议，费用不高于市场均价。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第六条 双方权利义务
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第六条  双方权利义务', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '6.1  甲方义务：', first_line_indent=0.74)
add_paragraph_cn(doc, '（1）提供施工场地、强弱电接口权限及必要协调支持；', first_line_indent=0.74)
add_paragraph_cn(doc, '（2）及时组织验收并按约付款；', first_line_indent=0.74)
add_paragraph_cn(doc, '（3）因甲方指定位置不符合安全规范导致的整改由甲方承担。', first_line_indent=0.74)
add_paragraph_cn(doc, '6.2  乙方义务：', first_line_indent=0.74)
add_paragraph_cn(doc, '（1）严格按国家电气与网络安全规范施工，确保施工质量；', first_line_indent=0.74)
add_paragraph_cn(doc, '（2）施工现场落实安全措施，承担施工期间自身人员伤亡及设备损坏责任；', first_line_indent=0.74)
add_paragraph_cn(doc, '（3）文明施工，竣工后清理现场垃圾。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第七条 违约责任
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第七条  违约责任', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '7.1  甲方逾期付款的，每逾期一日按应付未付款项的万分之三向乙方支付违约金；'
    '逾期超过 15 日，乙方有权暂停服务或解除合同。', first_line_indent=0.74)
add_paragraph_cn(doc, '7.2  乙方未按期完工或验收不合格经两次整改仍不达标的，甲方有权要求减免价款、'
    '重新施工或单方解除本合同，并要求赔偿直接损失。', first_line_indent=0.74)
add_paragraph_cn(doc, '7.3  任何一方违反本合同约定给对方造成损失的，应承担相应的赔偿责任。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第八条 不可抗力
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第八条  不可抗力', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '8.1  因不可抗力（包括但不限于自然灾害、战争、政府行为等）导致合同无法履行或需延期履行的，'
    '受影响方应及时通知对方并提供证明，双方协商顺延工期或解除合同，互不承担违约责任。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第九条 保密条款
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第九条  保密条款', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '9.1  甲乙双方应对在合同履行过程中知悉的对方商业秘密、技术资料及其他未公开信息予以保密，'
    '未经对方书面同意不得向第三方披露或用于本合同目的之外的用途。本保密义务不因合同终止而解除。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第十条 争议解决
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第十条  争议解决', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '10.1  因履行本合同发生的争议，双方应友好协商；协商不成的，'
    '任何一方均可向项目所在地人民法院提起诉讼。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 第十一条 附则
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '第十一条  附则', font_name='黑体', size=14, bold=True)
add_paragraph_cn(doc, '11.1  本合同一式 贰 份，甲乙双方各执 壹 份，具有同等法律效力。', first_line_indent=0.74)
add_paragraph_cn(doc, '11.2  本合同自双方法定代表人或授权代表签字并加盖公章（或合同专用章）之日起生效。', first_line_indent=0.74)
add_paragraph_cn(doc, '11.3  附件为本合同不可分割的组成部分，与正文具有同等效力。', first_line_indent=0.74)
add_paragraph_cn(doc, '11.4  本合同未尽事宜，经双方协商一致后，可签订补充协议。补充协议与本合同具有同等法律效力。', first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 附件一：报价单
# ═══════════════════════════════════════════════
add_paragraph_cn(doc, '', size=12)
add_paragraph_cn(doc, '附件一：报价单', font_name='黑体', size=14, bold=True, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph_cn(doc, '', size=6)

# 创建报价单表格
table = doc.add_table(rows=4, cols=6)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = 'Table Grid'

# 表头数据
headers = ['序号', '品名', '数量', '单位', '单价（元）', '小计（元）']
data_rows = [
    ['1', '交换机 H3C 八口千兆', '3', '台', '165', '495'],
    ['2', '人工费（含布线、安装、调试）', '10', '个', '350', '3,500'],
    ['', '合计', '', '', '', '3,995'],
]

# 设置列宽
col_widths = [Cm(1.5), Cm(5.5), Cm(1.5), Cm(1.5), Cm(2.0), Cm(2.0)]
for i, width in enumerate(col_widths):
    for row in table.rows:
        row.cells[i].width = width

# 填充表头
for i, header in enumerate(headers):
    cell = table.cell(0, i)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(header)
    set_cn_font(run, '宋体', 10.5, bold=True)
    run.font.color.rgb = RGBColor(255, 255, 255)
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    shade_cell(cell, '333333')
    set_cell_vertical_alignment(cell, 'center')

# 填充数据行
for row_idx, row_data in enumerate(data_rows):
    for col_idx, value in enumerate(row_data):
        cell = table.cell(row_idx + 1, col_idx)
        p = cell.paragraphs[0]
        # 金额列右对齐
        if col_idx >= 4:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        elif col_idx == 0:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if col_idx <= 3 else WD_ALIGN_PARAGRAPH.LEFT

        run = p.add_run(value)
        is_last = (row_idx == len(data_rows) - 1)  # 合计行
        set_cn_font(run, '宋体', 10.5, bold=is_last)
        p.paragraph_format.line_spacing = 1.2
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        set_cell_vertical_alignment(cell, 'center')

        # 合计行背景
        if is_last:
            shade_cell(cell, 'F2F2F2')

add_paragraph_cn(doc, '', size=6)
add_paragraph_cn(doc, '备注：以上报价为含税价，包含设备费、运输费、安装费、调试费及税费。', font_name='宋体', size=9, first_line_indent=0.74)

# ═══════════════════════════════════════════════
# 分页 - 签署页
# ═══════════════════════════════════════════════
doc.add_page_break()

add_paragraph_cn(doc, '', size=12)
add_paragraph_cn(doc, '（本页为签署页，无正文）', font_name='宋体', size=10.5, alignment=WD_ALIGN_PARAGRAPH.CENTER)
add_paragraph_cn(doc, '', size=24)

# 签署表格
sign_table = doc.add_table(rows=7, cols=2)
sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER

sign_data = [
    ('甲方（盖章）', '乙方（盖章）'),
    ('', '________________________'),
    ('', ''),
    ('授权代表签字', '授权代表签字'),
    ('__________________', '________________________'),
    ('', ''),
    ('日期：______年____月____日', '日期：______年____月____日'),
]

for row_idx, (left, right) in enumerate(sign_data):
    for col_idx, text in enumerate([left, right]):
        cell = sign_table.cell(row_idx, col_idx)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.8
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_cn_font(run, '宋体', 12, bold=(row_idx in [0, 3]))
        set_cell_vertical_alignment(cell, 'center')

        # Remove borders
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(f'<w:tcBorders {nsdecls("w")}></w:tcBorders>')
        tcPr.append(tcBorders)

sign_table.cell(0, 0).width = Cm(7)
sign_table.cell(0, 1).width = Cm(7)

# ═══════════════════════════════════════════════
# 添加页脚 - 页码
# ═══════════════════════════════════════════════
for section in doc.sections:
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add page number field
    run = p.add_run()
    fldChar1 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>')
    run._r.append(fldChar1)
    run2 = p.add_run()
    instrText = parse_xml(f'<w:instrText {nsdecls("w")} xml:space="preserve"> PAGE </w:instrText>')
    run2._r.append(instrText)
    run3 = p.add_run()
    fldChar2 = parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>')
    run3._r.append(fldChar2)

# ── 保存 ──
output_path = '/Users/yourname/Desktop/芜湖古城物业停车场系统和监控系统网络分离项目施工合同_商用版.docx'
doc.save(output_path)
print(f'✅ 合同已保存至: {output_path}')
